from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api.formatters import TextFormatter
from docx import Document
import os

app = FastAPI()

@app.get("/transcript-to-docx/")
def transcript_to_docx(video_id: str):
    try:
        # Get transcript
        transcript = YouTubeTranscriptApi.get_transcript(video_id)
        formatter = TextFormatter()
        transcript_text = formatter.format_transcript(transcript)
        
        # Create Word document
        doc = Document()
        doc.add_heading('YouTube Transcript', 0)
        doc.add_paragraph(transcript_text)

        # Save to file
        file_path = f"{video_id}.docx"
        doc.save(file_path)

        # Serve file
        return FileResponse(file_path, filename=file_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

